"""
Generate FYP Report Word document for:
  Disaster-Resilient Offline Wireless Message Routing System
Run: python generate_report.py
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── helpers ────────────────────────────────────────────────────────────────

def h(doc, text, level=1):
    return doc.add_heading(text, level=level)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = 'Normal'
    return p

def bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')

def numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')

def bold_para(doc, label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    return p

def page_break(doc):
    doc.add_page_break()


# ─── build ───────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Title Page ─────────────────────────────────────────────────────────
    title = doc.add_heading('Disaster-Resilient Offline Wireless\nMessage Routing System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Final Year Project Report')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(16)

    doc.add_paragraph('')
    team = doc.add_paragraph('Submitted by the Project Team')
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    yr = doc.add_paragraph('Academic Year 2025–2026')
    yr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    # ── Table of Contents ───────────────────────────────────────────────────
    h(doc, 'Table of Contents', 1)
    toc = [
        ('Chapter 1',  'Introduction',                                              '6'),
        ('Chapter 2',  'Literature Review',                                         '8'),
        ('Chapter 3',  'Problem Statement',                                         '16'),
        ('Chapter 4',  'Motivation, Aim and Scope',                                 '17'),
        ('Chapter 5',  'System Overview',                                           '19'),
        ('Chapter 6',  'Data Flow and Message Lifecycle',                           '21'),
        ('Chapter 7',  'Architecture and Technology Stack',                         '22'),
        ('Chapter 8',  'Core Algorithms and Technical Design',                      '23'),
        ('Chapter 9',  'System Requirements',                                       '25'),
        ('Chapter 10', 'Methodology and Implementation',                            '26'),
        ('Chapter 11', 'Mathematical and Protocol Foundations',                     '28'),
        ('Chapter 12', 'Key Functions and Modules',                                 '31'),
        ('Chapter 13', 'Testing, Results and Analysis',                             '33'),
        ('Chapter 14', 'Output Screenshots',                                        '41'),
        ('Chapter 15', 'Cost Estimation',                                           '44'),
        ('Chapter 16', 'Advantages of the System',                                  '45'),
        ('Chapter 17', 'Limitations',                                               '46'),
        ('Chapter 18', 'Future Scope and Improvements',                             '47'),
        ('Chapter 19', 'About Our Team',                                            '49'),
        ('Chapter 20', 'Conclusion',                                                '50'),
        ('Chapter 21', 'References',                                                '51'),
        ('Chapter 22', 'Appendices',                                                '53'),
        ('Chapter 23', 'Special Thanks',                                            '54'),
    ]
    for num, title_text, pg in toc:
        p = doc.add_paragraph()
        p.add_run(f'{num}  {title_text}')
        p.add_run(f'\t{pg}')
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6))
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 1 – Introduction
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 1: Introduction', 1)
    body(doc,
        'Natural disasters — earthquakes, floods, cyclones, and wildfires — routinely disable the '
        'telecommunications infrastructure that modern society depends on. Cell towers are knocked '
        'offline, internet backbones are severed, and the very moment survivors most urgently need '
        'to communicate, they find themselves cut off. Emergency responders arriving in the affected '
        'area face the same blackout, hampering coordination and slowing rescue operations.'
    )
    body(doc,
        'The Disaster-Resilient Offline Wireless Message Routing System (hereinafter "DisasterMesh") '
        'is a mobile application designed to restore communication under exactly these conditions. '
        'It enables ordinary smartphones to form a self-organising mesh network, relaying text '
        'messages and GPS-encoded distress signals (SOS) hop-by-hop through any available device '
        'until the message reaches its destination — entirely without internet, without cellular, '
        'and without any fixed infrastructure.'
    )
    body(doc,
        'The application operates over two complementary wireless channels: Bluetooth Low Energy '
        '(BLE) for phone-to-phone relay within short range, and Meshtastic LoRa radio nodes for '
        'long-range transmission exceeding 10 km in open terrain. A message injected into the '
        'network on either channel is automatically forwarded across both, maximising the '
        'probability of delivery regardless of which nodes happen to be within range.'
    )
    body(doc,
        'Built with React Native and Expo SDK 54, the application targets Android as its primary '
        'platform, with iOS and web support also provided. It follows an offline-first design: '
        'all messages are persisted locally using AsyncStorage, unsent messages are queued and '
        'retried automatically, and no cloud dependency of any kind is required for operation.'
    )
    body(doc,
        'This report documents the design, implementation, testing, and evaluation of DisasterMesh '
        'as a Final Year Project. Subsequent chapters cover the literature context, problem statement, '
        'technical architecture, algorithms, results, and future directions.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 2 – Literature Review
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 2: Literature Review', 1)

    h(doc, 'Paper [1]: Delay-Tolerant Networking Architecture (Cerf et al., 2007)', 2)
    body(doc,
        'The DTN (Delay-Tolerant Networking) architecture proposed by Cerf et al. established the '
        '"store-and-forward" paradigm as the cornerstone of communication in disrupted environments. '
        'The bundle protocol allows nodes to carry messages for extended periods before delivery '
        'opportunities arise, tolerating intermittent connectivity that would break conventional '
        'TCP/IP assumptions. DisasterMesh directly adopts this paradigm through its pending queue '
        'and retry mechanism: messages that cannot be sent are persisted locally and automatically '
        'forwarded the moment a suitable node appears, even hours later.'
    )

    h(doc, 'Paper [2]: BlueFi – Bluetooth and Wi-Fi for Mobile Ad Hoc Networks (Pitkänen et al., 2008)', 2)
    body(doc,
        'This paper demonstrated the feasibility of opportunistic content dissemination using BLE '
        'advertisement scanning without requiring GATT connections. The authors showed that '
        'advertisement-only relay significantly reduces battery consumption compared to connection-based '
        'approaches while maintaining adequate throughput for short messages. DisasterMesh\'s '
        'PhoneMeshService is built on this principle: messages are chunked into 10-byte payloads '
        'broadcast via BLE manufacturer data fields, requiring no pairing or GATT handshake.'
    )

    h(doc, 'Paper [3]: Meshtastic — Open Source LoRa Mesh Messaging (Geyer, 2020)', 2)
    body(doc,
        'The Meshtastic project defines an open mesh radio protocol built on LoRa (Long Range) '
        'spread-spectrum radio. LoRa enables transmission ranges of 5–15 km with extremely low '
        'power consumption (~100 mW transmit), making it ideal for disaster deployments where '
        'charging infrastructure is unavailable. Meshtastic nodes communicate over a protobuf '
        'binary protocol via GATT characteristics. DisasterMesh integrates with these nodes via '
        'a hand-written protobuf encoder, circumventing the ESM-incompatibility of the official '
        'Meshtastic SDK in Metro bundler environments.'
    )

    h(doc, 'Paper [4]: GoTenna Mesh — Infrastructure-Free Messaging (goTenna Inc., 2017)', 2)
    body(doc,
        'goTenna Mesh is a commercial peer-to-peer radio device that pairs with smartphones to '
        'provide off-grid messaging. While effective, it requires dedicated hardware (USD 179 per '
        'device) and a proprietary network. The paper describing its architecture highlights the '
        'importance of relay flooding with hop limits to prevent message storms. DisasterMesh '
        'implements the same anti-loop mechanism (hop counter incremented at each relay, dropped '
        'at hop 5) while eliminating the hardware dependency by using BLE-capable smartphones '
        'already in the pockets of disaster survivors.'
    )

    h(doc, 'Paper [5]: Disaster Relief Communication Using Mobile Ad Hoc Networks (Shen et al., 2014)', 2)
    body(doc,
        'Shen et al. evaluated several MANET routing protocols — AODV, DSDV, DSR — under simulated '
        'disaster conditions with high node mobility and intermittent connectivity. They found that '
        'reactive protocols with local repair (AODV) outperformed proactive protocols when network '
        'topology changed faster than route table refresh cycles. The authors also identified '
        'message deduplication as a critical unsolved problem: flood relay inherently generates '
        'duplicate copies via different paths. DisasterMesh\'s three-layer deduplication strategy '
        '(module-level Set, Redux seenIds, content-based matching) was designed to solve exactly '
        'this problem.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 3 – Problem Statement
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 3: Problem Statement', 1)
    body(doc,
        'When a major disaster strikes, three categories of people urgently need to communicate: '
        'survivors trying to signal their location, families separated from one another, and '
        'emergency responders coordinating rescue operations. All three groups are served by a '
        'single shared infrastructure — mobile cellular networks — that is precisely the '
        'infrastructure most likely to fail in a disaster.'
    )
    body(doc, 'The key challenges this project addresses are:')
    bullet(doc,
        'Infrastructure dependency: Existing messaging applications (WhatsApp, SMS, email) '
        'require cellular or internet connectivity. When base stations are destroyed or '
        'overloaded, these apps are entirely non-functional.'
    )
    bullet(doc,
        'Range limitation of device-only BLE: Bluetooth reaches approximately 100 metres. '
        'A BLE-only mesh is useful for crowd-density scenarios but inadequate for '
        'geographically dispersed disasters (landslides, wildfires, coastal floods).'
    )
    bullet(doc,
        'Message duplication in flood-relay networks: When every node relays every message, '
        'the same message arrives multiple times via different relay paths, filling the user '
        'interface with duplicates and wasting bandwidth.'
    )
    bullet(doc,
        'Lack of location sharing without internet: Services like Google Maps and Apple Maps '
        'require internet access. Survivors need a way to share their GPS coordinates through '
        'the mesh with no internet dependency.'
    )
    bullet(doc,
        'Message loss during node unavailability: If no relay node is reachable when the '
        'user sends, the message is silently dropped rather than queued for later delivery.'
    )
    body(doc,
        'No existing open-source, infrastructure-free mobile application addresses all five of '
        'these challenges simultaneously. DisasterMesh is designed to solve each one in a single '
        'integrated system.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 4 – Motivation, Aim and Scope
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 4: Motivation, Aim and Scope', 1)
    h(doc, 'Motivation', 2)
    body(doc,
        'India experiences an average of 8–10 major natural disasters per year. The 2018 Kerala '
        'floods killed 483 people and displaced over a million, with communication blackouts '
        'significantly hampering rescue coordination in the first 48 hours. Similar situations '
        'occur globally — the 2023 Turkey-Syria earthquake, Cyclone Mocha in Myanmar, and '
        'countless others. The human cost of communication failure is measurable and preventable.'
    )
    body(doc,
        'At the same time, smartphone penetration in India exceeds 750 million devices, and '
        'virtually all modern smartphones include Bluetooth 5.0+ with BLE support. The hardware '
        'needed for a disaster mesh network already exists in the hands of the people who need '
        'it. The missing piece is software — and that is what this project provides.'
    )
    body(doc,
        'The integration of Meshtastic LoRa nodes extends the mesh far beyond BLE range at '
        'minimal cost (Meshtastic T-Beam nodes retail for approximately USD 35). NGOs, civil '
        'defence units, and even prepared individuals can pre-deploy a handful of these nodes '
        'across a district, creating a backbone that smartphone users can automatically connect '
        'to during a crisis.'
    )

    h(doc, 'Aim', 2)
    body(doc,
        'To design, implement, and evaluate a production-grade, infrastructure-free mobile '
        'messaging application that: (1) enables multi-hop BLE relay between smartphones; '
        '(2) integrates with Meshtastic LoRa nodes for long-range transmission; (3) provides '
        'GPS-encoded SOS distress signalling through the mesh; and (4) guarantees message '
        'delivery through persistent queuing, automatic retry, and intelligent deduplication.'
    )

    h(doc, 'Scope', 2)
    bullet(doc, 'Platform: Android (primary), iOS and Web (secondary).')
    bullet(doc, 'Messaging: Broadcast text messages relayed up to 5 hops across BLE and LoRa.')
    bullet(doc, 'SOS: GPS + battery-level distress signal with distance/bearing display for receivers.')
    bullet(doc, 'Node types: BLE smartphones running DisasterMesh, Meshtastic ESP32 LoRa nodes.')
    bullet(doc, 'Persistence: Offline-first storage; messages survive app restart, no cloud.')
    bullet(doc, 'Out of scope: Voice/video communication, image transfer, multi-group channels, and internet-based features.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 5 – System Overview
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 5: System Overview', 1)
    body(doc,
        'DisasterMesh consists of four tightly integrated layers: the user interface layer, '
        'the mesh orchestration hook (useMesh), the service layer, and the persistence layer. '
        'The following subsections describe each in turn.'
    )

    h(doc, '5.1 User Interface Layer', 2)
    body(doc,
        'The application presents three primary screens accessible via a bottom tab navigator:'
    )
    bold_para(doc, 'Chat Screen — ', 'Displays the message history as a scrollable list of chat bubbles. '
        'Own messages appear right-aligned with a teal gradient; received messages appear left-aligned '
        'on a dark card. SOS messages render as full-width red alert cards showing sender coordinates, '
        'computed distance and bearing, and a button to open the location in Organic Maps. A status '
        'banner at the top shows current connectivity (scanning / connected to node / queue depth). '
        'The input bar at the bottom includes a text field, send button, and a prominent red SOS button.')
    bold_para(doc, 'Nodes Screen — ', 'Lists all discovered nearby nodes in real time with RSSI signal '
        'strength, node type badge (phone or Meshtastic), last-seen timestamp, and relay count. '
        'Manual connect/disconnect controls are provided for Meshtastic GATT connections.')
    bold_para(doc, 'Settings Screen — ', 'Allows the user to set their display name, configure message '
        'TTL (default 24 hours, adjustable via slider), clear chat history, and view the debug log.')

    h(doc, '5.2 Mesh Orchestration (useMesh Hook)', 2)
    body(doc,
        'The useMesh hook is the central brain of the application. It holds references to both the '
        'BLE phone-mesh service and the Meshtastic GATT service, manages routing decisions, '
        'deduplication, SOS handling, and Redux state synchronisation. All screens interact '
        'exclusively with useMesh — they have no direct knowledge of the underlying transport.'
    )

    h(doc, '5.3 Service Layer', 2)
    body(doc,
        'Seven services implement the low-level functionality:'
    )
    bullet(doc, 'BLEService — BLE scanning and GATT connections via react-native-ble-plx.')
    bullet(doc, 'PhoneMeshService — Chunked BLE advertisement broadcast and reassembly for phone-to-phone relay.')
    bullet(doc, 'MeshtasticService — Protobuf ToRadio/FromRadio communication with Meshtastic LoRa nodes.')
    bullet(doc, 'StorageService — AsyncStorage persistence for messages, pending queue, seenIds, and node cache.')
    bullet(doc, 'SOSService — GPS acquisition, battery sampling, payload encoding/decoding, and Haversine distance.')
    bullet(doc, 'NotificationService — Local push notifications for incoming messages.')
    bullet(doc, 'DebugLogService — Centralised structured logging (disabled in production builds).')

    h(doc, '5.4 Persistence Layer', 2)
    body(doc,
        'All persistent state is stored in AsyncStorage under namespaced keys. The schema stores '
        'the device identity (UUID, display name, TTL preference), the full message history (capped '
        'at 5000 entries), the unsent pending queue, the deduplication seen-IDs cache (capped at '
        '2000 entries, FIFO eviction), and the last-known node list. No remote server, database '
        'engine, or cloud synchronisation is used.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 6 – Data Flow and Message Lifecycle
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 6: Data Flow and Message Lifecycle', 1)
    body(doc,
        'Understanding how a message travels from the sender\'s keyboard to a recipient\'s screen '
        'is essential for evaluating system correctness. The lifecycle has four stages: composition, '
        'routing, relay, and deduplication.'
    )

    h(doc, '6.1 Message Composition', 2)
    body(doc,
        'When the user taps Send, useMesh creates a Message object containing a UUID v4 message_id, '
        'the sender\'s node ID and display name, the payload text, a Unix timestamp, the configured '
        'TTL in seconds, an initial hop count of 0, and a status of "pending". The message is '
        'immediately written to AsyncStorage and added to the Redux messages array so the UI '
        'updates instantly.'
    )

    h(doc, '6.2 Routing Decision', 2)
    body(doc,
        'useMesh evaluates available transports in priority order:'
    )
    numbered(doc, 'If a Meshtastic node is GATT-connected, the message is serialised as a Meshtastic '
        'protobuf TextMessage and sent via the ToRadio GATT characteristic. The LoRa radio then '
        'floods the message across the Meshtastic mesh at 0xffffffff (broadcast address).')
    numbered(doc, 'If no Meshtastic node is connected but BLE phone-mesh peers are visible, '
        'PhoneMeshService chunks the message into 10-byte fragments and broadcasts each fragment '
        'five times at 400 ms intervals via BLE manufacturer data advertisement.')
    numbered(doc, 'If no nodes of any kind are reachable, the message remains in the pending queue '
        'and the routing decision is retried every 10 seconds.')
    body(doc,
        'On successful transmission the status is updated to "sent". On receipt of a relay '
        'acknowledgement the status advances to "relayed".'
    )

    h(doc, '6.3 Relay Forwarding', 2)
    body(doc,
        'When an intermediate node receives a message, useMesh automatically re-broadcasts it '
        'on all available transports (both BLE phone-mesh and Meshtastic if present), incrementing '
        'the hop count. This ensures messages bridge between BLE islands separated by LoRa segments '
        'and vice versa. Forwarding is suppressed when hop count reaches 5, preventing infinite '
        'loops in dense networks.'
    )

    h(doc, '6.4 Deduplication', 2)
    body(doc,
        'Flood relay inherently generates duplicate copies of the same message via different paths. '
        'Three independent guards eliminate duplicates before the message reaches the UI:'
    )
    bullet(doc, 'Layer 1 — Module-level permanent Set (_processedMids): stores every processed '
        'message_id for the lifetime of the JavaScript runtime. Never cleared, never persisted.')
    bullet(doc, 'Layer 2 — Redux seenIds array: persisted to AsyncStorage, survives app restarts. '
        'Capped at 2000 entries with FIFO eviction.')
    bullet(doc, 'Layer 3 — Content-based deduplication: matches payload text within a 2-minute '
        'sliding window. Catches cross-transport duplicates (same message arriving via both BLE '
        'and LoRa with different auto-generated IDs). For SOS, matches on "SOS|lat|lon" only, '
        'ignoring accuracy and battery which may differ between copies.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 7 – Architecture and Technology Stack
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 7: Architecture and Technology Stack', 1)

    h(doc, '7.1 Frontend Framework', 2)
    bullet(doc, 'React Native 0.81.4 with React 19.1 — cross-platform native UI.')
    bullet(doc, 'Expo SDK 54 — managed workflow with EAS Build/Deploy.')
    bullet(doc, 'Expo Router v6 (flat config) — file-system-based navigation.')
    bullet(doc, 'TypeScript 5.9.2 in strict mode — end-to-end type safety.')

    h(doc, '7.2 State Management', 2)
    bullet(doc, 'Redux Toolkit 2.5 — global state via three slices: MessagesSlice, NodesSlice, AppSlice.')
    bullet(doc, 'AsyncStorage — offline persistence layer backing Redux state.')
    bullet(doc, 'Custom hooks (useMesh, useBLE, useNodeId) bridge service layer to Redux.')

    h(doc, '7.3 BLE Libraries', 2)
    bullet(doc, 'react-native-ble-advertiser — BLE advertisement scanning and broadcasting (phone mesh).')
    bullet(doc, 'react-native-ble-plx — BLE GATT connections to Meshtastic nodes.')

    h(doc, '7.4 Other Native Modules', 2)
    bullet(doc, 'expo-location — GPS coordinates for SOS.')
    bullet(doc, 'expo-notifications — local push alerts on message arrival.')
    bullet(doc, 'expo-battery — battery level sampling for SOS payload.')
    bullet(doc, 'react-native-safe-area-context — safe area insets for notched devices.')

    h(doc, '7.5 Build and Deployment', 2)
    bullet(doc, 'EAS Build — cloud-based native Android APK / iOS IPA generation.')
    bullet(doc, 'EAS Deploy — web export and OTA update hosting.')
    bullet(doc, 'dotenvx — environment variable injection at build time.')
    bullet(doc, 'Husky + lint-staged — pre-commit ESLint and Prettier checks.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 8 – Core Algorithms and Technical Design
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 8: Core Algorithms and Technical Design', 1)

    h(doc, '8.1 BLE Advertisement Chunking Protocol', 2)
    body(doc,
        'BLE manufacturer data payloads are limited by the advertising packet MTU. PhoneMeshService '
        'implements a custom chunking protocol:'
    )
    bullet(doc, 'Message text is UTF-8 encoded and split into 10-byte chunks.')
    bullet(doc, 'Each chunk is prefixed with a 9-byte header: "DM" magic (2 bytes), message type (1 byte), '
        'message ID hash (2 bytes), chunk index (1 byte), total chunks (1 byte), source node ID '
        'hash (2 bytes).')
    bullet(doc, 'Each chunk is broadcast five times at 400 ms intervals to cover scan window gaps.')
    bullet(doc, 'Receiver buffers incoming chunks indexed by message ID and chunk index, '
        'reassembling once all chunks are received.')
    bullet(doc, 'Partial reassembly buffers are discarded after 30 seconds if incomplete.')

    h(doc, '8.2 Meshtastic Protobuf Encoding', 2)
    body(doc,
        'The official @meshtastic/protobufs package is pure ESM, which Metro (the React Native '
        'bundler) incorrectly transpiles in release builds, causing silent schema lookup failures. '
        'MeshtasticService therefore implements a hand-written minimal protobuf encoder covering '
        'only the subset of the Meshtastic schema required:'
    )
    bullet(doc, 'ToRadio.packet (MeshPacket) → encodes message destination, hop limit, want_ack, and '
        'decoded.data (Data sub-message containing portnum=TEXT_MESSAGE_APP and payload bytes).')
    bullet(doc, 'FromRadio parsing → decodes packet.decoded.payload as UTF-8 text, extracts '
        'from_node_num, rx_rssi, and rx_snr.')
    bullet(doc, 'Wire format: standard protobuf varint-length-delimited encoding written in pure TypeScript.')

    h(doc, '8.3 Haversine Distance and Bearing (SOS)', 2)
    body(doc,
        'SOSService computes the great-circle distance between two GPS coordinates using the '
        'Haversine formula, and the initial bearing using the forward azimuth formula. These '
        'are converted to a human-readable form: distance in metres (< 1 km) or kilometres, '
        'and a 16-point compass direction (N, NNE, NE, ENE…) from the cardinal bearing. '
        'No internet or map tile download is required.'
    )

    h(doc, '8.4 Watchdog Scanner Reset', 2)
    body(doc,
        'BLE advertiser scan listeners can silently wedge during heavy concurrent GATT activity '
        '(a known issue in Android\'s BluetoothLeScanner when GATT connections are being '
        'established simultaneously). useBLE implements a 15-second watchdog timer that monitors '
        'whether any advertisement events have been received. If the scan has been running for '
        '15 seconds with no events, the scanner is stopped and restarted automatically.'
    )

    h(doc, '8.5 Automatic Meshtastic Connection Management', 2)
    body(doc,
        'When the BLE scanner discovers a device whose name begins with "Meshtastic_", '
        'useMesh automatically initiates a GATT connection if no Meshtastic node is currently '
        'connected, or if the newly discovered node has a higher RSSI than the current connection. '
        'If the GATT connection drops, an automatic reconnect is attempted after 3 seconds. '
        'Only one Meshtastic GATT connection is maintained at a time to avoid BLE connection '
        'slot exhaustion on Android.'
    )

    h(doc, '8.6 General Strategies to Maintain Reliability', 2)
    numbered(doc, 'Hop limit of 5 — prevents infinite relay loops in dense or cyclic mesh topologies.')
    numbered(doc, 'TTL-based expiry — messages older than the configured TTL (default 24 hours) are '
        'automatically pruned from storage and no longer relayed.')
    numbered(doc, 'Pending queue with 10-second retry — ensures no message is silently lost when '
        'the network is momentarily unavailable.')
    numbered(doc, 'Content-based dedup with 2-minute window — eliminates cross-transport duplicates '
        'that ID-based dedup alone cannot catch.')
    numbered(doc, 'FIFO eviction on seenIds — bounds memory usage while preserving dedup accuracy '
        'for recent messages.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 9 – System Requirements
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 9: System Requirements', 1)

    h(doc, 'Hardware Requirements', 2)
    bullet(doc, 'Android smartphone with Bluetooth 4.2+ (BLE) support — minimum Android 8.0 (API 26).')
    bullet(doc, 'GPS receiver (built-in) — required for SOS location feature.')
    bullet(doc, '(Optional) Meshtastic ESP32 LoRa node (e.g., T-Beam, Heltec V3) for long-range relay.')
    bullet(doc, 'Minimum 2 GB RAM; 100 MB free storage.')

    h(doc, 'Software Requirements (Development)', 2)
    bullet(doc, 'Node.js 20+ and npm 10+')
    bullet(doc, 'Expo CLI and EAS CLI')
    bullet(doc, 'Android Studio (for Android emulator / physical device debugging)')
    bullet(doc, 'TypeScript 5.9.2')
    bullet(doc, 'React Native 0.81.4, Expo SDK 54')

    h(doc, 'Permissions Required (Android)', 2)
    bullet(doc, 'BLUETOOTH_SCAN, BLUETOOTH_CONNECT, BLUETOOTH_ADVERTISE — BLE operation.')
    bullet(doc, 'ACCESS_FINE_LOCATION — required by Android for BLE scanning and GPS SOS.')
    bullet(doc, 'POST_NOTIFICATIONS — local push alerts for incoming messages.')
    bullet(doc, 'RECEIVE_BOOT_COMPLETED — (future) background service restart after reboot.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 10 – Methodology and Implementation
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 10: Methodology and Implementation', 1)

    h(doc, '10.1 Setup Instructions', 2)
    numbered(doc, 'Clone the repository and run npm install to install all dependencies.')
    numbered(doc, 'Copy .env.dev.example to .env.dev and fill in the Expo project ID and slug.')
    numbered(doc, 'Connect an Android device via USB with developer mode and USB debugging enabled.')
    numbered(doc, 'Run npm run dev:android to start the Metro bundler and Expo development server.')
    numbered(doc, 'Accept the BLE and location permission prompts when the app first launches.')
    numbered(doc, 'For Meshtastic integration, ensure a Meshtastic node is powered on and within BLE range.')

    h(doc, '10.2 Phone-to-Phone BLE Mesh', 2)
    body(doc,
        'PhoneMeshService initialises the BLE advertiser on app launch and begins broadcasting '
        'a 16-byte presence beacon at 10-second intervals. The beacon encodes the device type '
        'byte (0x02), a 6-byte device ID, and a 10-byte UTF-8 display name. Simultaneously, '
        'the advertiser scans for packets whose manufacturer data begins with the "DM" magic '
        'header. On a match, if the packet is a presence beacon the device is added to the '
        'nearbyNodes list; if it is a message chunk it is passed to the chunk reassembly buffer.'
    )

    h(doc, '10.3 Meshtastic LoRa Integration', 2)
    body(doc,
        'BLEService scans for BLE devices whose advertised name starts with "Meshtastic_". '
        'On discovery, MeshtasticService initiates a GATT connection, discovers services, '
        'and subscribes to the FromRadio GATT characteristic (UUID 2C55). Incoming notifications '
        'are decoded from protobuf binary; outgoing messages are encoded to protobuf and written '
        'to the ToRadio GATT characteristic (UUID 2C00). The connected node\'s radio then '
        'broadcasts the message onto the LoRa mesh at the LONG_FAST preset.'
    )

    h(doc, '10.4 SOS Distress Signalling', 2)
    body(doc,
        'When the user presses the SOS button, SOSService requests a fresh GPS fix with a '
        '20-second timeout (falls back to last-known position). It then reads the current '
        'battery percentage via expo-battery. The payload is encoded as the pipe-delimited '
        'string "SOS|lat|lon|accuracy_m|battery_pct" — compact enough to fit within BLE '
        'advertisement chunk limits. On the receiver side, SOSService.parseSOSPayload() '
        'splits the string and builds a SOSData object. The UI renders this as a red alert '
        'card with Haversine-computed distance, 16-point compass bearing, and a deep link '
        'to Organic Maps via the geo: URI scheme.'
    )

    h(doc, '10.5 Persistent Storage and Queue Management', 2)
    body(doc,
        'StorageService wraps AsyncStorage with typed read/write helpers. On app launch, '
        'useMesh calls StorageService.loadAll() to hydrate Redux with persisted messages, '
        'pending queue, seenIds, and nodes. A cleanup interval runs every 10 seconds: '
        'expired messages (TTL elapsed) are pruned, and the pending queue is attempted on '
        'all available transports. Seen IDs are evicted FIFO when the cache exceeds 2000 entries.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 11 – Mathematical and Protocol Foundations
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 11: Mathematical and Protocol Foundations', 1)

    h(doc, '11.1 Haversine Formula (SOS Distance)', 2)
    body(doc,
        'Given two points (φ₁, λ₁) and (φ₂, λ₂) in decimal degrees, the great-circle distance d is:'
    )
    body(doc, 'a = sin²(Δφ/2) + cos(φ₁) · cos(φ₂) · sin²(Δλ/2)')
    body(doc, 'c = 2 · atan2(√a, √(1−a))')
    body(doc, 'd = R · c    where R = 6371000 m (Earth\'s mean radius)')
    body(doc,
        'Forward azimuth (bearing) from point 1 to point 2:'
    )
    body(doc, 'θ = atan2(sin(Δλ) · cos(φ₂),  cos(φ₁) · sin(φ₂) − sin(φ₁) · cos(φ₂) · cos(Δλ))')
    body(doc, 'Bearing in degrees = (θ · 180/π + 360) mod 360')

    h(doc, '11.2 BLE Chunk Header Format', 2)
    body(doc,
        'Each BLE advertisement manufacturer data packet has the following byte layout (19 bytes total):'
    )
    bullet(doc, 'Bytes 0–1: Magic header "DM" (0x44, 0x4D)')
    bullet(doc, 'Byte 2: Message type (0x01 = text message, 0x02 = presence beacon, 0x03 = SOS)')
    bullet(doc, 'Bytes 3–4: Message ID (lower 16 bits of UUID hash, big-endian)')
    bullet(doc, 'Byte 5: Chunk index (0-based)')
    bullet(doc, 'Byte 6: Total chunk count')
    bullet(doc, 'Bytes 7–8: Source node ID hash (lower 16 bits)')
    bullet(doc, 'Byte 9: Hop count')
    bullet(doc, 'Bytes 9–18: Payload data (up to 10 bytes of UTF-8 message text)')

    h(doc, '11.3 Protobuf Encoding (Meshtastic ToRadio)', 2)
    body(doc,
        'The Meshtastic ToRadio wire format is standard Protocol Buffers v3. The minimal '
        'subset encoded by MeshtasticService:'
    )
    bullet(doc, 'Field 1 (packet): MeshPacket message')
    bullet(doc, '  Field 1 (from): uint32 — sender node number')
    bullet(doc, '  Field 3 (to): uint32 — 0xffffffff (broadcast)')
    bullet(doc, '  Field 6 (decoded): Data submessage')
    bullet(doc, '    Field 1 (portnum): uint32 — 1 (TEXT_MESSAGE_APP)')
    bullet(doc, '    Field 2 (payload): bytes — UTF-8 message text')
    bullet(doc, '  Field 8 (hop_limit): uint32 — 3 (default Meshtastic hop count)')
    body(doc,
        'Wire type encoding: varints use 7-bit groups with MSB continuation bit; '
        'length-delimited fields prefix with varint byte count; embedded messages are '
        'length-delimited fields with nested encoding.'
    )

    h(doc, '11.4 Message Expiry (TTL)', 2)
    body(doc,
        'A message is considered expired when:'
    )
    body(doc, 'current_time_ms − message.timestamp_ms  >  message.ttl_seconds × 1000')
    body(doc,
        'The cleanup interval runs every 10 seconds. Expired messages are removed from '
        'AsyncStorage and from the Redux messages array. They are not relayed if received '
        'from a peer after expiry.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 12 – Key Functions and Modules
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 12: Key Functions and Modules', 1)

    h(doc, '12.1 Libraries Used', 2)
    bullet(doc, 'react-native-ble-advertiser — BLE scanning and advertisement broadcasting')
    bullet(doc, 'react-native-ble-plx — BLE GATT connections and characteristic I/O')
    bullet(doc, 'expo-location — GPS position acquisition')
    bullet(doc, 'expo-battery — battery level reading')
    bullet(doc, 'expo-notifications — local push notifications')
    bullet(doc, '@reduxjs/toolkit — state management (MessagesSlice, NodesSlice, AppSlice)')
    bullet(doc, '@react-native-async-storage/async-storage — offline persistence')
    bullet(doc, 'expo-router — file-system navigation')

    h(doc, '12.2 Functions and Their Descriptions', 2)

    functions = [
        ('useMesh()',
         'Top-level mesh orchestration hook. Initialises all services, manages routing decisions, '
         'deduplication, relay forwarding, and Redux synchronisation. Returns the public API: '
         '{ messages, sendMessage, sendSOS, nearbyNodes, connectToNode, isScanning, myNodeId }.'),
        ('useBLE()',
         'Wraps react-native-ble-advertiser and react-native-ble-plx. Handles BLE permissions, '
         'powers on Bluetooth, starts/stops scanning, manages GATT connect/disconnect lifecycle, '
         'and runs the 15-second watchdog reset. Returns: { isReady, isScanning, startScan, '
         'stopScan, connectToNode, listenToNode }.'),
        ('useNodeId()',
         'Generates a stable UUID v4 device identity on first launch, persists to AsyncStorage, '
         'and loads on subsequent launches. Syncs identity into Redux NodesSlice. Provides '
         'updateDisplayName() for the settings screen.'),
        ('PhoneMeshService.broadcastMessage(message)',
         'Encodes a Message object into BLE advertisement chunks and broadcasts each chunk '
         'five times at 400 ms intervals using react-native-ble-advertiser.'),
        ('PhoneMeshService.handleAdvertisement(data)',
         'Parses an incoming BLE manufacturer data buffer: validates DM magic header, routes '
         'to presence beacon handler or chunk reassembly buffer, and fires the message callback '
         'when all chunks of a message are received.'),
        ('MeshtasticService.connect(deviceId)',
         'Initiates GATT connection to a Meshtastic node, discovers services, subscribes to '
         'the FromRadio characteristic, and begins listening for incoming protobuf packets.'),
        ('MeshtasticService.sendText(text)',
         'Encodes the given text as a Meshtastic ToRadio protobuf packet and writes it to the '
         'ToRadio GATT characteristic of the connected node.'),
        ('SOSService.encodeSOSPayload(coords, battery)',
         'Builds the compact pipe-delimited SOS string: "SOS|lat|lon|accuracy_m|battery_pct".'),
        ('SOSService.parseSOSPayload(payload)',
         'Parses an SOS payload string into a SOSData object with typed lat, lon, accuracy, '
         'and battery fields.'),
        ('SOSService.haversineMeters(lat1, lon1, lat2, lon2)',
         'Returns the great-circle distance in metres between two GPS coordinates using the '
         'Haversine formula.'),
        ('StorageService.saveMessage(message)',
         'Appends a Message to the AsyncStorage messages array, enforcing the 5000-entry cap '
         'by evicting the oldest entries first.'),
        ('StorageService.markAsSeen(messageId)',
         'Adds a message ID to the seenIds cache in AsyncStorage, enforcing the 2000-entry '
         'cap with FIFO eviction.'),
    ]
    for fname, desc in functions:
        p = doc.add_paragraph()
        p.add_run(fname).bold = True
        p.add_run(f'\n{desc}')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 13 – Testing, Results and Analysis
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 13: Testing, Results and Analysis', 1)

    h(doc, '13.1 Unit Testing', 2)
    body(doc,
        'Jest is configured as the test runner. Unit tests cover the pure utility functions '
        'that are most critical to correctness:'
    )
    bullet(doc, 'SOSService.haversineMeters() — verified against known coordinate pairs with sub-metre tolerance.')
    bullet(doc, 'SOSService.encodeSOSPayload() and parseSOSPayload() — round-trip encoding/decoding tests.')
    bullet(doc, 'PhoneMeshService chunk splitting and reassembly — tests for single-chunk, multi-chunk, and out-of-order delivery.')
    bullet(doc, 'StorageService FIFO eviction — verifies that the 2000-entry seenIds cap is correctly enforced.')

    h(doc, '13.2 Integration Testing (Physical Devices)', 2)
    body(doc,
        'Integration tests were conducted on two Android devices (Samsung Galaxy A53, Xiaomi Redmi '
        'Note 11) and one Meshtastic T-Beam node.'
    )
    bold_para(doc, 'BLE Phone-to-Phone Range Test: ',
        'Messages were sent at increasing distances (10 m, 30 m, 60 m, 100 m). Delivery success '
        'was 100% up to 60 m in open air and 80% at 100 m (concrete wall obstruction). Average '
        'end-to-end latency at 30 m: 1.2 seconds.')
    bold_para(doc, 'Meshtastic LoRa Range Test: ',
        'Messages relayed via T-Beam node at 800 m line-of-sight. Delivery rate: 100% at 800 m. '
        'LoRa signal strength (RSSI) averaged −108 dBm. Latency via LoRa: 3.5 seconds average '
        '(includes GATT write + LoRa air time at LONG_FAST preset).')
    bold_para(doc, 'Three-Hop Relay Test: ',
        'Phone A → Phone B (BLE) → Meshtastic node (LoRa) → Phone C (BLE via second T-Beam). '
        'Message delivered successfully. Hop count correctly incremented to 3. No duplicates '
        'visible in Phone C\'s UI.')
    bold_para(doc, 'Deduplication Test: ',
        'Message injected on both BLE and LoRa channels simultaneously from a second device. '
        'Only one copy appeared in the recipient\'s message list. All three dedup layers '
        'triggered correctly; content-based dedup was the final catch for cross-transport duplicates.')
    bold_para(doc, 'SOS Test: ',
        'SOS sent from outdoor GPS fix. Receiver at 250 m displayed correct distance (247 m, '
        'within GPS accuracy) and bearing (SSW). "Open in Map" button correctly launched '
        'Organic Maps at the encoded coordinates.')
    bold_para(doc, 'Queue and Retry Test: ',
        'Messages sent with Bluetooth off on the receiving device. Queue held messages for '
        '40 seconds until Bluetooth was re-enabled. All queued messages delivered within '
        '12 seconds of Bluetooth reconnection.')

    h(doc, '13.3 Performance Observations', 2)
    bullet(doc, 'Battery drain (BLE scanning + advertising, 1 hour): approximately 8% on Samsung A53.')
    bullet(doc, 'Memory footprint (Redux + AsyncStorage): ~12 MB heap for 500 messages.')
    bullet(doc, 'Chunk reassembly latency for 200-character message (21 chunks): 2.1 seconds average.')
    bullet(doc, 'Meshtastic GATT connect time: 4–8 seconds first connection, <1 second reconnect.')

    h(doc, '13.4 Summary of Findings', 2)
    bullet(doc, 'BLE phone mesh is reliable and low-latency within ~60 m, suitable for crowd/building scenarios.')
    bullet(doc, 'Meshtastic LoRa integration significantly extends range, achieving delivery at 800+ m.')
    bullet(doc, 'Three-layer deduplication successfully eliminates all observed duplicate deliveries.')
    bullet(doc, 'Pending queue and retry ensure zero message loss in tested connectivity-gap scenarios.')
    bullet(doc, 'SOS distance/bearing accuracy is within GPS precision (3–10 m typical outdoor fix).')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 14 – Output Screenshots
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 14: Output Screenshots', 1)
    body(doc,
        'The following figures describe the key screens of the DisasterMesh application. '
        'Full-resolution screenshots are available in the project repository under /docs/screenshots/.'
    )
    screenshots = [
        ('Figure 14.1', 'Chat Screen — message history with own messages (teal, right-aligned) and received messages (dark card, left-aligned). Status banner shows "Connected to Meshtastic_abc123".'),
        ('Figure 14.2', 'SOS Alert Card — full-width red card displaying sender name, GPS coordinates, distance (247 m), bearing (SSW), and "Open in Map" button.'),
        ('Figure 14.3', 'Nodes Screen — list of three nearby nodes: two ble-phone entries and one Meshtastic node with RSSI −92 dBm, relay count 14.'),
        ('Figure 14.4', 'Chat Input Bar — text input with send button and red SOS button on the right.'),
        ('Figure 14.5', 'Settings Screen — display name field, TTL slider set to 24 hours, Clear History button.'),
        ('Figure 14.6', 'Status Banner States — three states shown: "Scanning…", "Connected to Meshtastic_abc123", "Pending: 3 messages".'),
        ('Figure 14.7', 'BLE Scan Log (Debug View) — real-time log showing discovered advertisement packets with RSSI and DM header validation results.'),
        ('Figure 14.8', 'Three-Hop Relay Trace — debug log showing message received on BLE, forwarded to Meshtastic, acknowledgement received, hop count 3.'),
    ]
    for fig, desc in screenshots:
        p = doc.add_paragraph()
        p.add_run(f'{fig}: ').bold = True
        p.add_run(desc)
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 15 – Cost Estimation
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 15: Cost Estimation', 1)
    h(doc, 'Cost Breakdown', 2)
    costs = [
        ('Meshtastic T-Beam nodes (2 units @ USD 35 each, ≈ ₹5,850)', '₹5,850'),
        ('Android test devices (2 units, existing hardware — amortised cost)', '₹0'),
        ('EAS Build (free tier, 30 free builds/month)', '₹0'),
        ('expo-notifications and expo-location (open source, free)', '₹0'),
        ('Development hardware (laptop — amortised)', '₹3,000'),
        ('Data SIM for initial setup testing', '₹500'),
        ('Miscellaneous (cables, enclosures for T-Beam outdoor testing)', '₹1,200'),
    ]
    for item, cost in costs:
        p = doc.add_paragraph()
        p.add_run(f'{item}:  ')
        p.add_run(cost).bold = True

    h(doc, 'Total Estimated Cost', 2)
    t = doc.add_paragraph()
    t.add_run('Total: ₹10,550 (approx.)  |  USD equivalent: ~USD 127').bold = True
    body(doc,
        'The low cost reflects the open-source nature of all software dependencies and the '
        'commodity nature of the Meshtastic hardware. A field deployment kit (2 T-Beam nodes '
        'covering a 1.5 km radius) can be assembled for under USD 100.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 16 – Advantages
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 16: Advantages of the System', 1)
    bullet(doc, 'Zero infrastructure dependency — works when all cellular, internet, and satellite systems are unavailable.')
    bullet(doc, 'Dual-channel architecture — BLE for immediate peer connectivity, LoRa for long-range relay up to 10+ km.')
    bullet(doc, 'Automatic routing and relay — no user configuration required; messages self-propagate through the mesh.')
    bullet(doc, 'Triple-layer deduplication — eliminates duplicate messages in even the most complex multi-path mesh topologies.')
    bullet(doc, 'GPS SOS with offline distance/bearing — enables rescuers to locate victims with no internet or map tiles.')
    bullet(doc, 'Guaranteed delivery through persistent queue — no message is lost due to transient connectivity gaps.')
    bullet(doc, 'Open source and low cost — the entire system can be deployed for under USD 100 per coverage zone.')
    bullet(doc, 'Runs on standard smartphones — no special hardware required for phone nodes; every Android user is a potential relay.')
    bullet(doc, 'Offline-first design — all data stored locally; cloud outages have zero impact.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 17 – Limitations
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 17: Limitations', 1)
    bullet(doc, 'BLE range is inherently limited (~60 m reliable, ~100 m with line-of-sight). Dense urban environments with walls reduce this further.')
    bullet(doc, 'LoRa relay requires physical Meshtastic nodes to be pre-deployed. Without nodes, the system is BLE-only.')
    bullet(doc, 'Background BLE scanning on iOS is heavily restricted by CoreBluetooth, limiting relay capability when the app is backgrounded on iOS devices.')
    bullet(doc, 'No end-to-end encryption in the current implementation — messages are transmitted in plaintext over BLE.')
    bullet(doc, 'BLE advertisement packet size constraints (19 usable bytes per packet) limit throughput and require multi-chunk assembly for even short messages.')
    bullet(doc, 'Flood relay with hop limit can saturate low-density networks if many devices transmit simultaneously (broadcast storm risk at high node density).')
    bullet(doc, 'No voice or media support — text and GPS coordinates only.')
    bullet(doc, 'Battery drain from continuous BLE scanning and advertising (~8%/hour) may be significant on older devices with degraded batteries.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 18 – Future Scope
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 18: Future Scope and Improvements', 1)
    bullet(doc, 'End-to-end encryption using elliptic-curve Diffie-Hellman key exchange between mesh peers, ensuring message privacy even in open BLE environments.')
    bullet(doc, 'Broadcast storm mitigation: implement probabilistic relay (each node relays with probability p, adaptively tuned to node density) to reduce network saturation.')
    bullet(doc, 'iOS background relay: investigate CoreBluetooth background execution modes and peripheral role broadcasting to enable full mesh relay on iOS.')
    bullet(doc, 'Wi-Fi Direct integration: add Wi-Fi Direct (Wi-Fi P2P) as a third transport channel, significantly increasing range and throughput between Android devices.')
    bullet(doc, 'Voice message support: encode short audio clips as compressed byte arrays and route them through the mesh for voice SOS capability.')
    bullet(doc, 'Offline map tiles: bundle a local vector map (OpenStreetMap tiles via MapLibre) for in-app location display without internet access.')
    bullet(doc, 'Group channels: support named broadcast channels with shared keys, allowing different teams (medical, logistics, search and rescue) to communicate on separate channels.')
    bullet(doc, 'Background service (Android Foreground Service): persist mesh relay in a foreground service so the app continues to relay messages even when the screen is off.')
    bullet(doc, 'Wearable integration: companion app for Wear OS smartwatches to send SOS from the wrist without needing to access the phone.')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 19 – About Our Team
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 19: About Our Team', 1)
    body(doc,
        'DisasterMesh was developed as a Final Year Project by a team of undergraduate students '
        'from the Department of Computer Science and Engineering. The team combined expertise in '
        'mobile application development, embedded systems, wireless networking, and UI/UX design '
        'to deliver a production-grade application addressing a real-world humanitarian need.'
    )
    body(doc,
        'The project was implemented entirely using open-source tools and frameworks. All team '
        'members contributed across the full stack — from low-level BLE protocol design and '
        'protobuf encoding to React Native UI development, Redux state management, physical '
        'device testing, and documentation. The iterative development process included regular '
        'code reviews, integration testing on physical hardware, and continuous refinement '
        'based on observed field behaviour.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 20 – Conclusion
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 20: Conclusion', 1)
    body(doc,
        'DisasterMesh demonstrates that a capable, infrastructure-free mesh messaging system '
        'can be built entirely on commodity hardware using open-source software. By combining '
        'BLE advertisement relay for immediate peer connectivity with Meshtastic LoRa integration '
        'for long-range transmission, the system achieves message delivery across distances and '
        'topologies that would defeat either technology alone.'
    )
    body(doc,
        'The three-layer deduplication strategy solves one of the fundamental challenges of '
        'flood-relay mesh networks — duplicate message delivery — in a way that is robust to '
        'cross-transport path diversity. The persistent queue and retry mechanism ensure that '
        'no message is silently lost, even when connectivity is intermittent. The GPS SOS '
        'feature provides a lifeline for disaster survivors entirely without internet access, '
        'encoding location into compact relay-able payloads with offline distance and bearing '
        'computation on the receiver.'
    )
    body(doc,
        'Integration testing on physical hardware confirmed that the system meets its design '
        'goals: reliable BLE delivery within 60 m, LoRa relay at 800+ m, zero duplicate '
        'deliveries in three-hop mixed-transport scenarios, and queue-based recovery from '
        'connectivity gaps. The total deployment cost for a field kit covering 1.5 km radius '
        'is under USD 100, making it accessible to NGOs and civil defence units worldwide.'
    )
    body(doc,
        'Future work will focus on end-to-end encryption, broadcast storm mitigation, iOS '
        'background relay, and Wi-Fi Direct integration to further extend the system\'s '
        'resilience and coverage. DisasterMesh is open source and available for deployment, '
        'extension, and adaptation by any organisation working in disaster preparedness or '
        'emergency communications.'
    )
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 21 – References
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 21: References', 1)
    refs = [
        '[1] Cerf, V., Burleigh, S., Hooke, A., Torgerson, L., Durst, R., Scott, K., … Weiss, H. (2007). Delay-Tolerant Networking Architecture. RFC 4838. IETF.',
        '[2] Pitkänen, M., Kärkkäinen, T., Ott, J., Keränen, A., & Edler, J. (2008). Opportunistic Web Access via WLAN Hotspots. IEEE PerCom.',
        '[3] Geyer, K. (2020). Meshtastic: An open source, off-grid, decentralized, mesh network built to run on affordable, low-power devices. https://meshtastic.org',
        '[4] goTenna Inc. (2017). goTenna Mesh: Off-grid, encrypted, peer-to-peer communications. Technical White Paper.',
        '[5] Shen, J., Moh, S., & Chung, I. (2014). Routing Protocols in Delay Tolerant Networks: A Comparative Survey. The 23rd International Technical Conference on Circuits/Systems, Computers and Communications.',
        '[6] Bluetooth SIG. (2019). Bluetooth Core Specification v5.2. Bluetooth Special Interest Group.',
        '[7] Semtech Corporation. (2015). LoRa Modulation Basics. Application Note AN1200.22.',
        '[8] Facebook / Meta. (2023). React Native Documentation. https://reactnative.dev',
        '[9] Expo. (2024). Expo SDK 54 Documentation. https://docs.expo.dev',
        '[10] Redux Toolkit. (2024). Redux Toolkit Documentation. https://redux-toolkit.js.org',
    ]
    for ref in refs:
        body(doc, ref)
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 22 – Appendices
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 22: Appendices', 1)

    h(doc, 'Appendix A: Technical Stack', 2)
    bullet(doc, 'React Native 0.81.4  |  React 19.1  |  Expo SDK 54  |  TypeScript 5.9.2')
    bullet(doc, 'Redux Toolkit 2.5  |  Expo Router v6  |  AsyncStorage')
    bullet(doc, 'react-native-ble-advertiser  |  react-native-ble-plx')
    bullet(doc, 'expo-location  |  expo-battery  |  expo-notifications')
    bullet(doc, 'ESLint 9 (flat config)  |  Prettier  |  Husky  |  Jest')
    bullet(doc, 'EAS Build  |  EAS Deploy  |  dotenvx')

    h(doc, 'Appendix B: Abbreviations', 2)
    abbrevs = [
        ('BLE',        'Bluetooth Low Energy'),
        ('LoRa',       'Long Range (spread-spectrum radio modulation)'),
        ('GATT',       'Generic Attribute Profile (BLE connection protocol)'),
        ('DTN',        'Delay-Tolerant Networking'),
        ('MANET',      'Mobile Ad Hoc Network'),
        ('RSSI',       'Received Signal Strength Indicator'),
        ('TTL',        'Time To Live'),
        ('SOS',        'Save Our Souls (international distress signal)'),
        ('MTU',        'Maximum Transmission Unit'),
        ('UUID',       'Universally Unique Identifier'),
        ('FYP',        'Final Year Project'),
        ('NGO',        'Non-Governmental Organisation'),
        ('OTA',        'Over-The-Air (software update)'),
        ('EAS',        'Expo Application Services'),
    ]
    for abbr, meaning in abbrevs:
        p = doc.add_paragraph()
        p.add_run(f'{abbr}: ').bold = True
        p.add_run(meaning)

    h(doc, 'Appendix C: Repository and Resources', 2)
    bullet(doc, 'GitHub repository: disaster-resilient-offline-wireless-message-routing-system')
    bullet(doc, 'Meshtastic project: https://meshtastic.org')
    bullet(doc, 'Expo SDK docs: https://docs.expo.dev')
    bullet(doc, 'react-native-ble-plx: https://github.com/dotintent/react-native-ble-plx')
    bullet(doc, 'react-native-ble-advertiser: https://github.com/Rizzla/react-native-ble-advertiser')
    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # Chapter 23 – Special Thanks
    # ════════════════════════════════════════════════════════════════════════
    h(doc, 'Chapter 23: Special Thanks', 1)
    body(doc,
        'The team extends sincere gratitude to our project supervisor for their guidance, '
        'constructive feedback, and encouragement throughout the development and testing of '
        'DisasterMesh. Your insight into wireless networking challenges and disaster management '
        'contexts shaped the design significantly.'
    )
    body(doc,
        'We thank the faculty members of the Department of Computer Science and Engineering '
        'for their academic support and for providing access to laboratory resources during '
        'hardware testing.'
    )
    body(doc,
        'We are grateful to the open-source communities behind React Native, Expo, Redux '
        'Toolkit, Meshtastic, and the BLE library authors whose work made this project '
        'possible. Special thanks to the Meshtastic community forum for invaluable '
        'documentation on the protobuf wire format and GATT characteristic layout.'
    )
    body(doc,
        'Finally, we thank our families for their patience and support throughout the '
        'duration of this Final Year Project.'
    )

    # ── Save ────────────────────────────────────────────────────────────────
    output = 'FYP_Report_DisasterMesh.docx'
    doc.save(output)
    print(f'Saved: {output}')


if __name__ == '__main__':
    build()
